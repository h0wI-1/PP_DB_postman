"""
Генератор отчёта по лабораторной работе 4 в формате DOCX.
Стиль: Times New Roman 14pt, A4, поля L=3/R=1.5/T=2/B=2 см.
Заголовки разделов: жирный, по центру.
Основной текст: по ширине.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── КОНСТАНТЫ ────────────────────────────────────────────────────────────────
FONT      = "Times New Roman"
BODY_SIZE = Pt(14)

# ─── ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ ──────────────────────────────────────────────────

def set_font(run, bold=False, italic=False, size=BODY_SIZE):
    run.font.name  = FONT
    run.font.size  = size
    run.font.bold  = bold
    run.font.italic = italic
    # Принудительно выставляем шрифт для кириллических символов
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'),      FONT)
    rFonts.set(qn('w:hAnsi'),      FONT)
    rFonts.set(qn('w:cs'),         FONT)
    rFonts.set(qn('w:eastAsia'),   FONT)
    # Удаляем старый rFonts, если был
    old = rPr.find(qn('w:rFonts'))
    if old is not None:
        rPr.remove(old)
    rPr.insert(0, rFonts)


def add_paragraph(doc, text="", bold=False, italic=False,
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_cm=None,
                  space_before_pt=None, space_after_pt=None):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before_pt) if space_before_pt else None
    pf.space_after  = Pt(space_after_pt)  if space_after_pt  else Pt(0)
    if first_line_cm is not None:
        pf.first_line_indent = Cm(first_line_cm)
    if text:
        run = p.add_run(text)
        set_font(run, bold=bold, italic=italic)
    return p


def heading(doc, text, number=None):
    """Заголовок раздела: жирный, по центру, номер — опционально."""
    label = f"{number}. {text}" if number else text
    add_paragraph(doc, label, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                  space_before_pt=12, space_after_pt=6)


def subheading(doc, text, number=None):
    """Подзаголовок (5.1, 5.2...) — жирный, по левому краю."""
    label = f"{number}  {text}" if number else text
    add_paragraph(doc, label, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
                  space_before_pt=6, space_after_pt=3)


def body(doc, text, first_line=1.25):
    """Обычный текстовый абзац: отступ первой строки 1.25 см."""
    add_paragraph(doc, text, first_line_cm=first_line, space_after_pt=0)


def code_block(doc, text):
    """Кодовый блок: Courier New 11pt, серый фон, без отступа."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    # Серый фон
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'F2F2F2')
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name  = "Courier New"
    run.font.size  = Pt(10)
    run.font.bold  = False
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Courier New')
    rFonts.set(qn('w:hAnsi'), 'Courier New')
    rFonts.set(qn('w:cs'),    'Courier New')
    old = rPr.find(qn('w:rFonts'))
    if old is not None: rPr.remove(old)
    rPr.insert(0, rFonts)
    return p


def img_placeholder(doc, fig_num, caption_text):
    """Плейсхолдер для скриншота и подпись под ним."""
    # Серая рамка-заглушка
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(0)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'D9D9D9')
    pPr.append(shd)
    run = p.add_run(f"[ СКРИНШОТ — Рисунок {fig_num} ]")
    set_font(run, italic=True, size=Pt(12))
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    # Подпись
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(3)
    cap.paragraph_format.space_after  = Pt(6)
    r = cap.add_run(f"Рисунок {fig_num} – {caption_text}")
    set_font(r, italic=False)


def page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx.oxml.ns.qn and __import__('docx').enum.text.WD_BREAK.PAGE)


# ─── СОЗДАНИЕ ДОКУМЕНТА ───────────────────────────────────────────────────────
import docx

doc = Document()

# Параметры страницы (A4, поля как в образце)
section = doc.sections[0]
section.page_width   = Cm(21.0)
section.page_height  = Cm(29.7)
section.left_margin  = Cm(3.0)
section.right_margin = Cm(1.5)
section.top_margin   = Cm(2.0)
section.bottom_margin = Cm(2.0)

# Устанавливаем Times New Roman 14pt в стиль Normal
normal_style = doc.styles['Normal']
normal_style.font.name = FONT
normal_style.font.size = BODY_SIZE
# XML-уровень для кириллицы
from lxml import etree
rPr_el = normal_style.element.find(qn('w:rPr'))
if rPr_el is None:
    rPr_el = OxmlElement('w:rPr')
    normal_style.element.append(rPr_el)
rFonts_el = OxmlElement('w:rFonts')
rFonts_el.set(qn('w:ascii'),    FONT)
rFonts_el.set(qn('w:hAnsi'),    FONT)
rFonts_el.set(qn('w:cs'),       FONT)
rFonts_el.set(qn('w:eastAsia'), FONT)
old = rPr_el.find(qn('w:rFonts'))
if old is not None: rPr_el.remove(old)
rPr_el.insert(0, rFonts_el)

# ─── ТИТУЛЬНЫЙ ЛИСТ ───────────────────────────────────────────────────────────

def title_para(doc, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(14)):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    if text:
        r = p.add_run(text)
        set_font(r, bold=bold, size=size)
    return p

title_para(doc, "МИНИСТЕРСТВО ЦИФРОВОГО РАЗВИТИЯ, СВЯЗИ И МАССОВЫХ")
title_para(doc, "КОММУНИКАЦИЙ РОССИЙСКОЙ ФЕДЕРАЦИИ")
title_para(doc, "Ордена Трудового Красного Знамени федеральное государственное "
                "бюджетное образовательное учреждение высшего образования")
title_para(doc, '«Московский технический университет связи и информатики»')
title_para(doc, "")
title_para(doc, "")
title_para(doc, "Кафедра «Системное программирование»")
title_para(doc, "")
title_para(doc, "")
title_para(doc, "Отчёт по лабораторной работе", bold=True)
title_para(doc, "по дисциплине «Программная инженерия»")
title_para(doc, "на тему «Тестирование базы данных интернет-магазина через Postman»")
title_para(doc, "")
title_para(doc, "")
# Правый блок
p_exec = title_para(doc, "                                        Выполнил: студент группы ________",
                    align=WD_ALIGN_PARAGRAPH.LEFT)
title_para(doc, "                                        Солотонов А.Э.",
           align=WD_ALIGN_PARAGRAPH.LEFT)
title_para(doc, "")
title_para(doc, "                                        Проверил: ________________________",
           align=WD_ALIGN_PARAGRAPH.LEFT)
title_para(doc, "")
title_para(doc, "")
title_para(doc, "")
title_para(doc, "Москва 2026")

# Разрыв страницы
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(0)
from docx.oxml import OxmlElement as OE
run = p.add_run()
br = OE('w:br')
br.set(qn('w:type'), 'page')
run._r.append(br)

# ─── РАЗДЕЛ 1 ─────────────────────────────────────────────────────────────────
heading(doc, "Цель работы", 1)

body(doc, "Целью данной лабораторной работы является получение практических навыков "
     "в проектировании реляционных баз данных, подключении базы данных к REST API "
     "и тестировании запросов через клиент Postman. В ходе работы необходимо показать "
     "выполнение единичных запросов, запросов с параметрами, стандартных операций "
     "создания, изменения и удаления записей, а также реализовать авторизацию по "
     "токену и написать автотесты.")
body(doc, "В качестве предметной области выбран интернет-магазин техники. База данных "
     "включает четыре таблицы: клиенты, товары, статусы заказов и сами заказы.")

# ─── РАЗДЕЛ 2 ─────────────────────────────────────────────────────────────────
heading(doc, "Используемые инструменты", 2)

body(doc, "В работе использован следующий набор инструментов:")
body(doc, "PostgreSQL 14 — реляционная СУБД, в которой хранится база данных.")
body(doc, "PostgREST — надстройка над PostgreSQL, автоматически публикующая таблицы "
     "и хранимые функции базы данных в виде REST API без написания дополнительного кода. "
     "Данный инструмент выбран, поскольку он даёт минимальный по объёму конфигурационный "
     "файл и при этом поддерживает весь необходимый набор операций: GET, POST, PATCH, "
     "DELETE, фильтрацию через параметры URL и авторизацию по JWT.")
body(doc, "Postman — HTTP-клиент, через который выполняется тестирование всех "
     "запросов к API. Используется также для создания коллекции, работы с "
     "переменными окружения и запуска автотестов.")
body(doc, "DBeaver — графическая среда для управления базой данных, применялась для "
     "просмотра результатов SQL-скриптов и построения ER-диаграммы.")
body(doc, "Расширение pgjwt для PostgreSQL — используется для подписи JWT-токенов "
     "внутри хранимой функции авторизации.")

# ─── РАЗДЕЛ 3 ─────────────────────────────────────────────────────────────────
heading(doc, "Создание базы данных", 3)

body(doc, "База данных shop_db включает четыре взаимосвязанные таблицы.")
body(doc, "Таблица clients хранит информацию о клиентах: идентификатор, имя, "
     "уникальный адрес электронной почты, телефон, город и дату регистрации.")
body(doc, "Таблица items содержит данные о товарах: идентификатор, наименование, "
     "категорию, цену и количество на складе.")
body(doc, "Таблица order_status — справочник статусов заказа. Содержит пять записей: "
     "«Новый», «Оплачен», «Отправлен», «Доставлен», «Отменён».")
body(doc, "Таблица orders связывает клиента и товар через внешние ключи, хранит "
     "статус заказа, количество товара, итоговую сумму и дату создания.")
body(doc, "ER-диаграмма базы данных приведена на рисунке 1.")

img_placeholder(doc, 1, "Логическая модель базы данных интернет-магазина")

body(doc, "SQL-скрипт создания таблиц:")

for line in [
    "CREATE TABLE clients (",
    "    id             SERIAL PRIMARY KEY,",
    "    name           VARCHAR(100) NOT NULL,",
    "    email          VARCHAR(120) NOT NULL UNIQUE,",
    "    phone          VARCHAR(20),",
    "    city           VARCHAR(80),",
    "    registered_at  TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP",
    ");",
    "",
    "CREATE TABLE items (",
    "    id        SERIAL PRIMARY KEY,",
    "    name      VARCHAR(120)   NOT NULL,",
    "    category  VARCHAR(60)    NOT NULL,",
    "    price     NUMERIC(10, 2) NOT NULL CHECK (price >= 0),",
    "    stock     INTEGER        NOT NULL DEFAULT 0",
    ");",
    "",
    "CREATE TABLE order_status (",
    "    id     SERIAL PRIMARY KEY,",
    "    code   VARCHAR(20)  NOT NULL UNIQUE,",
    "    title  VARCHAR(60)  NOT NULL",
    ");",
    "",
    "CREATE TABLE orders (",
    "    id          SERIAL PRIMARY KEY,",
    "    client_id   INTEGER NOT NULL REFERENCES clients(id),",
    "    item_id     INTEGER NOT NULL REFERENCES items(id),",
    "    status_id   INTEGER NOT NULL REFERENCES order_status(id),",
    "    quantity    INTEGER NOT NULL DEFAULT 1,",
    "    total       NUMERIC(12, 2) NOT NULL,",
    "    created_at  TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP",
    ");",
]:
    code_block(doc, line)

body(doc, "После создания таблицы заполнены тестовыми данными: 6 клиентов, 8 товаров, "
     "5 статусов и 6 заказов. Данные подобраны так, чтобы запросы с фильтрацией "
     "(по цене, имени) сразу возвращали осмысленный результат.")

img_placeholder(doc, 2, "Заполнение таблицы clients тестовыми данными")
img_placeholder(doc, 3, "Заполнение таблицы items тестовыми данными")

# ─── РАЗДЕЛ 4 ─────────────────────────────────────────────────────────────────
heading(doc, "Подключение базы данных к API", 4)

body(doc, "Для обращения к базе данных через HTTP-протокол используется PostgREST. "
     "Настройки задаются в конфигурационном файле postgrest.conf:")

for line in [
    'db-uri        = "postgres://postgres:postgres@localhost:5432/shop_db"',
    'db-schema     = "public"',
    'db-anon-role  = "web_anon"',
    'jwt-secret    = "super_secret_key_change_me_please_32chars"',
    'server-port   = 3000',
]:
    code_block(doc, line)

body(doc, "Параметр db-uri задаёт строку подключения к базе данных. Параметр "
     "db-schema указывает схему, таблицы которой будут опубликованы как ресурсы "
     "REST API. Параметр db-anon-role определяет роль, под которой выполняются "
     "запросы без токена авторизации. Параметр jwt-secret — секрет, которым "
     "подписываются и проверяются JWT-токены. Параметр server-port задаёт порт, "
     "на котором PostgREST принимает HTTP-запросы.")
body(doc, "PostgREST запускается командой:")
code_block(doc, "postgrest config/postgrest.conf")
body(doc, "После запуска в терминале появляется строка «Listening on port 3000», "
     "что свидетельствует о готовности сервера к приёму запросов. "
     "Рисунки 4 и 5 демонстрируют конфигурацию и результат запуска.")

img_placeholder(doc, 4, "Конфигурационный файл PostgREST")
img_placeholder(doc, 5, "Запуск сервера PostgREST")

# ─── РАЗДЕЛ 5 ─────────────────────────────────────────────────────────────────
heading(doc, "Тестирование базы данных через Postman", 5)

subheading(doc, "Единичные GET-запросы", "5.1")

body(doc, "Первый этап тестирования — единичные запросы на получение всех записей "
     "из каждой таблицы. В Postman создаётся запрос с методом GET и адресом "
     "http://localhost:3000/clients. Сервер возвращает массив JSON из 6 объектов, "
     "каждый из которых соответствует одному клиенту. Аналогичным образом "
     "выполняется запрос GET /items — возвращается список из 8 товаров.")

img_placeholder(doc, 6, "GET-запрос на получение всех клиентов")
img_placeholder(doc, 7, "GET-запрос на получение всех товаров")

subheading(doc, "GET-запросы с параметрами", "5.2")

body(doc, "В PostgREST параметры фильтрации передаются в строке URL после знака «?». "
     "Формат параметра: имя_поля=оператор.значение. В работе использованы следующие "
     "виды параметров:")
body(doc, "GET /clients?id=eq.1 — выбор клиента по идентификатору. Оператор eq "
     "соответствует условию равенства (WHERE id = 1).")
body(doc, "GET /items?price=gt.40000 — товары дороже 40 000 рублей. Оператор gt "
     "соответствует условию «больше чем» (WHERE price > 40000).")
body(doc, "GET /clients?name=like.A* — клиенты, имя которых начинается на букву "
     "«A». Оператор like с символом * работает аналогично LIKE 'A%' в SQL.")
body(doc, "GET /clients?select=name — выборка только одного поля. В ответе "
     "возвращается массив объектов вида {\"name\": \"...\"}, остальные столбцы "
     "в выдаче отсутствуют.")
body(doc, "Все перечисленные запросы выполнены в Postman, результаты приведены "
     "на рисунках 8–11.")

img_placeholder(doc, 8,  "GET-запрос с параметром id=eq.1")
img_placeholder(doc, 9,  "Фильтр товаров по цене (price=gt.40000)")
img_placeholder(doc, 10, "Фильтр клиентов по шаблону имени (name=like.A*)")
img_placeholder(doc, 11, "Выборка отдельного поля (select=name)")

subheading(doc, "Стандартные запросы POST, PATCH, DELETE", "5.3")

body(doc, "В задании к лабораторной работе указан метод PULL. Поскольку "
     "в стандарте HTTP метода с таким именем не существует, а ближайший по "
     "семантике метод частичного обновления ресурса — PATCH, "
     "в данной работе используется именно он. PostgREST поддерживает PATCH "
     "в полном объёме.")
body(doc, "POST — добавление новой записи. Метод POST, адрес /clients, тело запроса "
     "в формате JSON:")

for line in [
    "{",
    '    "name":  "Test User",',
    '    "email": "test.user@example.com",',
    '    "phone": "+7-900-000-00-00",',
    '    "city":  "Тестоград"',
    "}",
]:
    code_block(doc, line)

body(doc, "В заголовке указано Prefer: return=representation — это заставляет "
     "PostgREST вернуть созданную запись вместе со сгенерированным идентификатором. "
     "Сервер возвращает статус 201 Created. После добавления выполняется проверочный "
     "GET-запрос, подтверждающий появление записи в базе.")

img_placeholder(doc, 12, "Добавление нового клиента (POST)")
img_placeholder(doc, 13, "Проверка добавления записи через GET")

body(doc, "PATCH — изменение записи. Метод PATCH, адрес /clients?id=eq.{id}, "
     "тело с изменяемыми полями:")
code_block(doc, '{ "phone": "+7-900-999-99-99", "city": "Москва" }')
body(doc, "Сервер обновляет только переданные поля и возвращает изменённую строку. "
     "Повторный GET подтверждает, что новые значения сохранились в базе.")

img_placeholder(doc, 14, "Изменение записи через PATCH")

body(doc, "DELETE — удаление записи. Метод DELETE, адрес /clients?id=eq.{id}. "
     "Сервер возвращает статус 200. Для проверки результата выполняется повторный "
     "GET с тем же параметром; в ответ приходит пустой массив [], что подтверждает "
     "отсутствие записи в базе данных.")

img_placeholder(doc, 15, "Удаление записи (DELETE)")
img_placeholder(doc, 16, "Проверка удаления записи (пустой ответ)")

# ─── РАЗДЕЛ 6 ─────────────────────────────────────────────────────────────────
heading(doc, "Создание коллекции запросов и её запуск", 6)

body(doc, "Все отлаженные запросы объединены в коллекцию Postman «Shop API». "
     "Для удобства коллекция разделена на четыре папки: «01. Авторизация», "
     "«02. GET-запросы без параметров», «03. GET-запросы с параметрами» "
     "и «04. CRUD».")
body(doc, "Чтобы запросы не зависели от жёстко прописанных адресов и значений, "
     "создан environment «Shop API Local» с переменными: base_url, token, "
     "new_client_id, login_email, login_password. В самих запросах вместо "
     "конкретных значений используются подстановки {{base_url}}, {{token}} и т. д.")
body(doc, "Для последовательного прогона коллекции используется встроенный в "
     "Postman Collection Runner: открывается вкладка Runner, выбирается "
     "коллекция «Shop API», нажимается кнопка Run. Запросы выполняются в "
     "порядке расположения в дереве коллекции, что важно для блока CRUD: "
     "запрос POST сначала создаёт запись и сохраняет её id в переменную, "
     "а затем PATCH, DELETE и проверочные GET используют эту переменную "
     "через {{new_client_id}}.")

img_placeholder(doc, 17, "Коллекция запросов «Shop API»")
img_placeholder(doc, 18, "Запуск коллекции через Collection Runner")

# ─── РАЗДЕЛ 7 ─────────────────────────────────────────────────────────────────
heading(doc, "Безопасность БД и авторизация по токену", 7)

body(doc, "По умолчанию PostgREST выполняет все запросы под анонимной ролью "
     "(параметр db-anon-role = web_anon). В данной работе этой роли выдан "
     "только SELECT на все четыре таблицы. Добавление, изменение и удаление "
     "данных доступны лишь роли web_user. Чтобы PostgREST выполнил запрос "
     "от имени web_user, необходим валидный JWT-токен.")
body(doc, "Схема авторизации реализована следующим образом. В базе данных "
     "создана таблица users с электронными адресами и хэшами паролей "
     "(функция crypt() из расширения pgcrypto). Хранимая функция "
     "public.login(email, password) проверяет правильность пароля и, "
     "если он верен, формирует JWT-токен, подписанный секретом jwt-secret. "
     "В теле токена указана роль web_user и время истечения (1 час).")
body(doc, "Клиент (Postman) отправляет запрос POST /rpc/login с телом:")
code_block(doc, '{ "email": "admin@shop.local", "password": "password123" }')
body(doc, "В ответ возвращается JSON с полем token. PostgREST проверяет подпись "
     "токена с помощью секрета из конфигурационного файла, извлекает из него "
     "роль и выполняет запрос уже от имени web_user.")
body(doc, "В Postman авторизация настроена на уровне коллекции: тип "
     "Authorization → Bearer Token, значение — {{token}}. В тестах запроса "
     "POST /rpc/login есть строка pm.environment.set('token', body.token), "
     "которая автоматически сохраняет полученный токен в переменную окружения, "
     "после чего он используется во всех последующих запросах.")

img_placeholder(doc, 19, "Получение JWT-токена (POST /rpc/login)")
img_placeholder(doc, 20, "Использование Bearer Token в Postman")

# ─── РАЗДЕЛ 8 ─────────────────────────────────────────────────────────────────
heading(doc, "Автотесты в Postman", 8)

body(doc, "Для каждого запроса коллекции во вкладке Tests написан набор "
     "проверок на языке JavaScript. Используется встроенный в Postman фреймворк "
     "Chai (через объект pm). Ниже приведён типовой набор проверок для "
     "GET-запроса:")

for line in [
    "pm.test('Status 200', function () {",
    "    pm.response.to.have.status(200);",
    "});",
    "pm.test('Response time < 1000ms', function () {",
    "    pm.expect(pm.response.responseTime).to.be.below(1000);",
    "});",
    "const body = pm.response.json();",
    "pm.test('Ответ — непустой массив', function () {",
    "    pm.expect(body).to.be.an('array').with.length.above(0);",
    "});",
    "pm.test('У клиента есть ключевые поля', function () {",
    "    pm.expect(body[0]).to.have.all.keys(",
    "        'id', 'name', 'email', 'phone', 'city', 'registered_at'",
    "    );",
    "});",
]:
    code_block(doc, line)

body(doc, "Для запроса POST дополнительно проверяется код 201 и наличие "
     "поля id в ответе; после этого значение id сохраняется в переменную "
     "окружения для использования в последующих запросах:")

for line in [
    "pm.test('Status 201', function () {",
    "    pm.response.to.have.status(201);",
    "});",
    "pm.environment.set('new_client_id', body[0].id);",
]:
    code_block(doc, line)

body(doc, "Для запроса PATCH проверяется, что возвращённая запись содержит "
     "обновлённое значение изменённого поля. Для запроса DELETE проверяется "
     "код 200, а следующий проверочный GET должен вернуть пустой массив:")

for line in [
    "pm.test('Массив пустой — записи больше нет', function () {",
    "    pm.expect(pm.response.json()).to.be.an('array').with.lengthOf(0);",
    "});",
]:
    code_block(doc, line)

body(doc, "При запуске коллекции через Collection Runner все проверки "
     "выполняются последовательно. В итоговой сводке отображается "
     "общее количество пройденных и упавших тестов.")

img_placeholder(doc, 21, "Автотесты во вкладке Tests")
img_placeholder(doc, 22, "Результаты прохождения автотестов")

# ─── РАЗДЕЛ 9 ─────────────────────────────────────────────────────────────────
heading(doc, "Вывод", 9)

body(doc, "В ходе выполнения лабораторной работы была спроектирована и создана "
     "учебная база данных интернет-магазина техники из четырёх взаимосвязанных "
     "таблиц, настроено подключение к ней через REST API с помощью PostgREST, "
     "а все основные сценарии работы с базой данных проверены через Postman.")
body(doc, "Рассмотрены единичные GET-запросы, запросы с параметрами фильтрации "
     "(равенство, сравнение по значению, поиск по шаблону, выборка отдельных "
     "полей), а также стандартные операции добавления, изменения и удаления "
     "записей. После каждой операции правильность результата подтверждалась "
     "повторным GET-запросом.")
body(doc, "Отдельно реализован раздел по безопасности: ограничены права "
     "анонимной роли, создана функция авторизации, возвращающая JWT-токен, "
     "настроена авторизация в Postman через заголовок Authorization: Bearer. "
     "Все запросы объединены в коллекцию и покрыты автотестами, которые "
     "позволяют одним запуском убедиться в работоспособности всей цепочки "
     "взаимодействия с базой данных.")
body(doc, "В результате получен рабочий учебный стенд, демонстрирующий полный "
     "цикл: от создания базы данных до автоматизированного тестирования "
     "REST API с авторизацией по токену.")

# ─── СОХРАНЕНИЕ ───────────────────────────────────────────────────────────────
out_path = "/Users/howis/Downloads/ПИ лаба 4/Отчёт_ПИ_лаба_4.docx"
doc.save(out_path)
print(f"Файл сохранён: {out_path}")
