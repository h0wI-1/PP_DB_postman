"""
Вставляет PNG-скриншоты в DOCX вместо серых плейсхолдеров.
Каждый плейсхолдер имеет вид '[ СКРИНШОТ — Рисунок N ]'.
После него стоит абзац с подписью 'Рисунок N – ...'.
"""

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import copy

DOCX_IN  = Path("/Users/howis/Downloads/ПИ лаба 4/Отчёт_ПИ_лаба_4.docx")
DOCX_OUT = Path("/Users/howis/Downloads/ПИ лаба 4/Отчёт_ПИ_лаба_4_final.docx")
SHOTS    = Path("/Users/howis/Downloads/ПИ лаба 4/screenshots")

# Карта: номер рисунка -> путь к PNG
import os, re
png_map = {}
for f in SHOTS.glob("fig*.png"):
    m = re.match(r"fig(\d+)_", f.name)
    if m:
        png_map[int(m.group(1))] = str(f)

print("Найдены PNG:", sorted(png_map.keys()))

doc = Document(str(DOCX_IN))
paragraphs = doc.paragraphs

i = 0
replaced = 0
while i < len(paragraphs):
    p = paragraphs[i]
    t = p.text
    m = re.search(r"СКРИНШОТ.*?Рисунок\s+(\d+)", t)
    if m:
        fig_num = int(m.group(1))
        png = png_map.get(fig_num)
        if not png:
            print(f"  [!] PNG для рис.{fig_num} не найден, пропускаю")
            i += 1
            continue

        # Очищаем плейсхолдер-абзац и вставляем картинку
        # Убираем shading (серый фон)
        pPr = p._p.find(qn('w:pPr'))
        if pPr is not None:
            shd = pPr.find(qn('w:shd'))
            if shd is not None:
                pPr.remove(shd)

        # Удаляем все runs из абзаца
        for r in p.runs:
            r._r.getparent().remove(r._r)

        # Выравнивание по центру
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(6)
        pf.space_after  = Pt(0)

        # Добавляем run с картинкой
        run = p.add_run()
        # Вставляем изображение, ширина = 14 см (влезает в поля)
        try:
            run.add_picture(png, width=Cm(14))
            print(f"  ✓ Рисунок {fig_num} вставлен")
            replaced += 1
        except Exception as e:
            run.text = f"[ошибка вставки рис.{fig_num}: {e}]"
            print(f"  [!] Ошибка рис.{fig_num}: {e}")

    i += 1

print(f"\nВставлено {replaced} из {len(png_map)} скриншотов.")
doc.save(str(DOCX_OUT))
print(f"Сохранён: {DOCX_OUT}")
