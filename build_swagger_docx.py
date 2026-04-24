"""
Генератор отчёта по Swagger в формате DOCX.
Стиль: Times New Roman 14pt, A4, поля L=3/R=1.5/T=2/B=2 см.
Скриншоты вставляются напрямую из screenshots/figSWNN_*.png.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import re

FONT      = "Times New Roman"
BODY_SIZE = Pt(14)
SHOTS     = Path("/Users/howis/Downloads/ПИ лаба 4/screenshots")
OUT_PATH  = "/Users/howis/Downloads/ПИ лаба 4/Отчёт_ПИ_лаба_4_Swagger.docx"

# ─── PNG-поиск ────────────────────────────────────────────────────────────────

png_map = {}
for f in SHOTS.glob("figSW*.png"):
    m = re.match(r"figSW(\d+)_", f.name)
    if m:
        png_map[int(m.group(1))] = str(f)
print("Swagger PNG найдены:", sorted(png_map.keys()))

# ─── ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ ──────────────────────────────────────────────────

def set_font(run, bold=False, italic=False, size=BODY_SIZE):
    run.font.name   = FONT
    run.font.size   = size
    run.font.bold   = bold
    run.font.italic = italic
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'),    FONT)
    rFonts.set(qn('w:hAnsi'),    FONT)
    rFonts.set(qn('w:cs'),       FONT)
    rFonts.set(qn('w:eastAsia'), FONT)
    old = rPr.find(qn('w:rFonts'))
    if old is not None:
        rPr.remove(old)
    rPr.insert(0, rFonts)


def add_paragraph(doc, text="", bold=False, italic=False,
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_cm=None,
                  space_before_pt=None, space_after_pt=None):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before_pt) if space_before_pt else None
    pf.space_after  = Pt(space_after_pt)  if space_after_pt  else Pt(0)
    if first_line_cm is not None:
        pf.first_line_indent = Cm(first_line_cm)
    if text:
        run = p.add_run(text)
        set_font(run, bold=bold, italic=italic)
    return p


def heading(doc, text, number=None):
    label = f"{number}. {text}" if number else text
    add_paragraph(doc, label, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                  space_before_pt=12, space_after_pt=6)


def subheading(doc, text, number=None):
    label = f"{number}  {text}" if number else text
    add_paragraph(doc, label, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
                  space_before_pt=6, space_after_pt=3)


def body(doc, text, first_line=1.25):
    add_paragraph(doc, text, first_line_cm=first_line, space_after_pt=0)


def code_block(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'F2F2F2')
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(10)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Courier New')
    rFonts.set(qn('w:hAnsi'), 'Courier New')
    rFonts.set(qn('w:cs'),    'Courier New')
    old = rPr.find(qn('w:rFonts'))
    if old is not None:
        rPr.remove(old)
    rPr.insert(0, rFonts)
    return p


def embed_image(doc, fig_num, caption_text, width_cm=14):
    png = png_map.get(fig_num)
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(6)
    p_img.paragraph_format.space_after  = Pt(0)
    run = p_img.add_run()
    if png:
        try:
            run.add_picture(png, width=Cm(width_cm))
            print(f"  ✓ Рисунок SW{fig_num:02d} вставлен")
        except Exception as e:
            run.text = f"[ошибка рис.SW{fig_num}: {e}]"
            print(f"  [!] Ошибка рис.SW{fig_num}: {e}")
    else:
        run.text = f"[PNG figSW{fig_num:02d} не найден]"
        print(f"  [!] PNG для рис.SW{fig_num} не найден")

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(3)
    cap.paragraph_format.space_after  = Pt(6)
    r = cap.add_run(f"Рисунок {fig_num} – {caption_text}")
    set_font(r, italic=False, size=Pt(12))


def page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)


# ─── СОЗДАНИЕ ДОКУМЕНТА ───────────────────────────────────────────────────────

doc = Document()

section = doc.sections[0]
section.page_width    = Cm(21.0)
section.page_height   = Cm(29.7)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(1.5)
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)

normal_style = doc.styles['Normal']
normal_style.font.name = FONT
normal_style.font.size = BODY_SIZE
from lxml import etree
rPr_el = normal_style.element.find(qn('w:rPr'))
if rPr_el is None:
    rPr_el = OxmlElement('w:rPr')
    normal_style.element.append(rPr_el)
rFonts_el = OxmlElement('w:rFonts')
rFonts_el.set(qn('w:ascii'),    FONT)
rFonts_el.set(qn('w:hAnsi'),    FONT)
rFonts_el.set(qn('w:cs'),       FONT)
rFonts_el.set(qn('w:eastAsia'), FONT)
old = rPr_el.find(qn('w:rFonts'))
if old is not None:
    rPr_el.remove(old)
rPr_el.insert(0, rFonts_el)

# ─── ТИТУЛЬНЫЙ ЛИСТ ───────────────────────────────────────────────────────────

def title_para(doc, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(14)):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    if text:
        r = p.add_run(text)
        set_font(r, bold=bold, size=size)
    return p

title_para(doc, "МИНИСТЕРСТВО ЦИФРОВОГО РАЗВИТИЯ, СВЯЗИ И МАССОВЫХ")
title_para(doc, "КОММУНИКАЦИЙ РОССИЙСКОЙ ФЕДЕРАЦИИ")
title_para(doc, "Ордена Трудового Красного Знамени федеральное государственное "
                "бюджетное образовательное учреждение высшего образования")
title_para(doc, '«Московский технический университет связи и информатики»')
title_para(doc, "")
title_para(doc, "")
title_para(doc, "Кафедра «Системное программирование»")
title_para(doc, "")
title_para(doc, "")
title_para(doc, "Отчёт по лабораторной работе №4 (дополнение)", bold=True)
title_para(doc, "по дисциплине «Программная инженерия»")
title_para(doc, "на тему «Документирование REST API с помощью Swagger / OpenAPI»")
title_para(doc, "")
title_para(doc, "")
title_para(doc, "                                        Выполнил: студент группы ________",
           align=WD_ALIGN_PARAGRAPH.LEFT)
title_para(doc, "                                        Солотонов А.Э.",
           align=WD_ALIGN_PARAGRAPH.LEFT)
title_para(doc, "")
title_para(doc, "                                        Проверил: ________________________",
           align=WD_ALIGN_PARAGRAPH.LEFT)
title_para(doc, "")
title_para(doc, "")
title_para(doc, "")
title_para(doc, "Москва 2026")

page_break(doc)

# ─── РАЗДЕЛ 1 ─────────────────────────────────────────────────────────────────
heading(doc, "Цель работы", 1)

body(doc, "Целью данного раздела лабораторной работы является освоение инструментария "
     "для документирования REST API на основе стандарта OpenAPI 3.0. В ходе работы "
     "разработана спецификация API интернет-магазина техники в формате YAML, после чего "
     "она была визуализирована через Swagger UI — интерактивную веб-страницу, которая "
     "позволяет просматривать документацию и отправлять запросы прямо из браузера.")

# ─── РАЗДЕЛ 2 ─────────────────────────────────────────────────────────────────
heading(doc, "Используемые инструменты", 2)

body(doc, "В работе использован следующий набор инструментов.")
body(doc, "OpenAPI 3.0.3 — стандарт описания REST API, разработанный Open API Initiative. "
     "Спецификация записывается в формате YAML или JSON и описывает все конечные точки API, "
     "параметры, тела запросов, ответы и схемы безопасности.")
body(doc, "Swagger UI — библиотека с открытым исходным кодом (версия 5), которая читает "
     "OpenAPI-файл и строит по нему интерактивную HTML-страницу. Не требует сборки: "
     "достаточно одного HTML-файла, подключающего библиотеку через CDN.")
body(doc, "PostgreSQL + PostgREST — стек API, задокументированный в спецификации. "
     "Swagger используется как слой документации поверх уже работающего API.")
body(doc, "Браузер — для открытия Swagger UI; дополнительных программ устанавливать не нужно.")

# ─── РАЗДЕЛ 3 ─────────────────────────────────────────────────────────────────
heading(doc, "Структура спецификации OpenAPI", 3)

body(doc, "Файл docs/swagger.yaml содержит полное описание API интернет-магазина и "
     "состоит из нескольких логических блоков.")

subheading(doc, "Заголовок и сервер", "3.1")

body(doc, "Блок info задаёт название, версию и контактные данные API. Блок servers "
     "указывает базовый URL: http://localhost:3000 — это порт, на котором работает PostgREST.")

for line in [
    "openapi: 3.0.3",
    "info:",
    '  title: Shop API — Интернет-магазин техники',
    '  version: "1.0.0"',
    "  contact:",
    "    email: artemsolotonov@gmail.com",
    "servers:",
    "  - url: http://localhost:3000",
    '    description: Локальный сервер PostgREST',
]:
    code_block(doc, line)

subheading(doc, "Теги", "3.2")

body(doc, "Все эндпоинты разбиты на пять групп: Auth, Clients, Items, OrderStatus, Orders. "
     "Теги используются для визуальной группировки в интерфейсе Swagger UI — каждая "
     "группа отображается как отдельный раскрываемый раздел.")

subheading(doc, "Схемы (components/schemas)", "3.3")

body(doc, "В блоке components описаны переиспользуемые схемы данных. Для каждой таблицы "
     "определено три схемы: полный объект из БД (включает id и дату создания), тело "
     "POST-запроса (без id и служебных полей) и тело PATCH-запроса (все поля необязательные). "
     "Дополнительно описаны схемы LoginRequest, LoginResponse и Error.")

subheading(doc, "Параметры (components/parameters)", "3.4")

body(doc, "Часто используемые параметры запросов вынесены в компонентные объекты: "
     "SelectParam, OrderParam, LimitParam, OffsetParam, PreferHeader. Это позволяет "
     "переиспользовать их в каждом эндпоинте через ссылку $ref, не дублируя описание.")

subheading(doc, "Схема безопасности", "3.5")

body(doc, "Используется схема BearerAuth типа http со схемой bearer. Токен передаётся "
     "в заголовке Authorization: Bearer <token>. В Swagger UI это реализуется через "
     "кнопку Authorize в верхней части страницы.")

for line in [
    "components:",
    "  securitySchemes:",
    "    BearerAuth:",
    "      type: http",
    "      scheme: bearer",
    "      bearerFormat: JWT",
]:
    code_block(doc, line)

subheading(doc, "Пути (paths)", "3.6")

body(doc, "Для каждого URL (/clients, /items, /order_status, /orders, /rpc/login) "
     "описаны все поддерживаемые методы: GET, POST, PATCH, DELETE. Каждый метод "
     "содержит summary и description с примерами запросов, список параметров со "
     "ссылками на схемы, requestBody для методов, принимающих тело, а также "
     "responses с кодами 200, 201, 204, 401, 409 и примерами JSON. Методы, не "
     "требующие авторизации (GET-запросы, /rpc/login), явно помечены security: [].")

embed_image(doc, 1, "Файл swagger.yaml в редакторе: общая структура документа.")
embed_image(doc, 2, "Блок components/schemas с описанием схем Client и ClientInput.")

# ─── РАЗДЕЛ 4 ─────────────────────────────────────────────────────────────────
heading(doc, "Запуск Swagger UI", 4)

body(doc, "Для отображения документации подготовлен файл docs/index.html. Он "
     "подключает библиотеку Swagger UI через CDN и передаёт ей путь к спецификации.")

for line in [
    "<script>",
    "  SwaggerUIBundle({",
    '    url: "./swagger.yaml",',
    '    dom_id: "#swagger-ui",',
    "    presets: [SwaggerUIBundle.presets.apis,",
    "              SwaggerUIBundle.SwaggerUIStandalonePreset],",
    '    layout: "BaseLayout",',
    "    deepLinking: true,",
    "    tryItOutEnabled: true",
    "  });",
    "</script>",
]:
    code_block(doc, line)

body(doc, "Параметр tryItOutEnabled: true активирует кнопку Try it out для всех "
     "операций сразу, без необходимости нажимать её вручную в каждом разделе.")
body(doc, "Для запуска достаточно запустить PostgREST и открыть docs/index.html "
     "через локальный HTTP-сервер, чтобы браузер мог загрузить файл swagger.yaml "
     "по относительному пути. Например:")

for line in [
    "# Запустить PostgREST",
    "postgrest config/postgrest.conf",
    "",
    "# В отдельном терминале запустить HTTP-сервер в папке docs/",
    'cd "path/to/ПИ лаба 4/docs"',
    "python3 -m http.server 8080",
]:
    code_block(doc, line)

body(doc, "После этого открыть в браузере: http://localhost:8080")

embed_image(doc, 3, "Запуск HTTP-сервера в терминале.")
embed_image(doc, 4, "Главная страница Swagger UI в браузере: заголовок, пять групп "
            "операций (Auth, Clients, Items, OrderStatus, Orders), кнопка Authorize.")

# ─── РАЗДЕЛ 5 ─────────────────────────────────────────────────────────────────
heading(doc, "Интерфейс Swagger UI", 5)

body(doc, "После открытия страницы отображается следующая структура. Вверху — "
     "название API, версия, описание архитектуры и базовый URL сервера. Ниже — "
     "секция Servers с выбором адреса (http://localhost:3000). Кнопка Authorize "
     "служит для ввода Bearer Token.")
body(doc, "Пять разделов (тегов) сворачиваются и разворачиваются кликом. Внутри "
     "каждого раздела — список операций, цвет рамки указывает на HTTP-метод: "
     "зелёный — GET, синий — POST, оранжевый — PATCH, красный — DELETE.")
body(doc, "При раскрытии операции видны параметры, схема тела запроса, примеры "
     "ответов и коды статусов.")

embed_image(doc, 5, "Раздел Clients с развёрнутым GET /clients: параметры фильтрации, "
            "примеры ответов.")
embed_image(doc, 6, "Примеры ответов (200) для GET /clients.")

# ─── РАЗДЕЛ 6 ─────────────────────────────────────────────────────────────────
heading(doc, "Авторизация в Swagger UI", 6)

body(doc, "Для выполнения защищённых запросов (POST, PATCH, DELETE) через кнопку "
     "Try it out необходимо сначала получить токен, а затем ввести его в интерфейс "
     "Swagger UI.")

subheading(doc, "Шаг 1 — получить токен", "")

body(doc, "Развернуть раздел Auth, операция POST /rpc/login. Нажать Try it out. "
     "В поле Request body ввести:")
for line in ["{", '    "email": "admin@shop.local",', '    "password": "password123"', "}"]:
    code_block(doc, line)
body(doc, "Нажать Execute. В блоке Responses → Response body скопировать значение поля token.")

subheading(doc, "Шаг 2 — авторизоваться", "")

body(doc, "В верхней части страницы нажать кнопку Authorize (иконка замка). В "
     "открывшемся окне в поле BearerAuth вставить скопированный токен. Нажать "
     "Authorize, затем Close.")
body(doc, "После этого все операции, помеченные иконкой замка, будут выполняться "
     "с заголовком Authorization: Bearer <token> автоматически. В интерфейсе "
     "замок рядом с операцией «закроется», сигнализируя об активной авторизации.")

embed_image(doc, 7, "Запрос POST /rpc/login через Try it out: тело запроса и ответ с токеном.")
embed_image(doc, 8, "Диалог Authorize: ввод Bearer Token.")
embed_image(doc, 9, "Иконка закрытого замка рядом с POST /clients после авторизации.")

# ─── РАЗДЕЛ 7 ─────────────────────────────────────────────────────────────────
heading(doc, "Тестирование запросов через Try it out", 7)

subheading(doc, "GET-запросы", "7.1")

body(doc, "GET /clients — развернуть операцию, нажать Execute. В поле Response body "
     "выводится массив JSON с клиентами, код ответа — 200.")
body(doc, "GET /clients с параметрами — заполнить поле id значением eq.1 и нажать "
     "Execute. В ответе — один объект с соответствующим клиентом.")
body(doc, "GET /items?price=gt.40000 — заполнить поле price значением gt.40000. "
     "В ответе — товары с ценой выше 40 000 ₽.")

subheading(doc, "POST — добавление клиента", "7.2")

body(doc, "Развернуть операцию POST /clients. В поле Request body ввести:")

for line in [
    "{",
    '    "name": "Test User",',
    '    "email": "test.user@example.com",',
    '    "phone": "+7-900-000-00-00",',
    '    "city": "Тестоград"',
    "}",
]:
    code_block(doc, line)

body(doc, "В заголовке Prefer указать return=representation. Нажать Execute. "
     "Swagger UI формирует полный HTTP-запрос с заголовком авторизации и отправляет "
     "его на http://localhost:3000/clients. В ответе — код 201 Created и JSON с "
     "созданной записью, включая сгенерированный id.")

subheading(doc, "PATCH — изменение записи", "7.3")

body(doc, "Развернуть операцию PATCH /clients. В параметре id указать eq.<id нового "
     "клиента>. В теле запроса:")
for line in ["{", '    "phone": "+7-900-999-99-99",', '    "city": "Москва"', "}"]:
    code_block(doc, line)
body(doc, "Нажать Execute. Код ответа — 200, в теле — обновлённая запись.")

subheading(doc, "DELETE — удаление записи", "7.4")

body(doc, "Развернуть операцию DELETE /clients. В параметре id указать eq.<id нового "
     "клиента>. Нажать Execute. Код ответа — 200 (или 204). Для проверки повторно "
     "выполнить GET /clients?id=eq.<id> — в ответе придёт пустой массив [].")

embed_image(doc, 10, "Try it out: GET /clients — ответ 200, массив клиентов.")
embed_image(doc, 11, "Try it out: GET /clients с параметром id=eq.1.")
embed_image(doc, 12, "Try it out: GET /items с параметром price=gt.40000.")
embed_image(doc, 13, "Try it out: POST /clients — заполненное тело запроса.")
embed_image(doc, 14, "Try it out: POST /clients — ответ 201, созданная запись с id.")
embed_image(doc, 15, "Try it out: PATCH /clients — параметр id и тело запроса.")
embed_image(doc, 16, "Try it out: DELETE /clients — ответ 200.")
embed_image(doc, 17, "Try it out: GET /clients после DELETE — пустой массив [].")

# ─── РАЗДЕЛ 8 ─────────────────────────────────────────────────────────────────
heading(doc, "Документирование схем", 8)

body(doc, "В нижней части страницы Swagger UI расположен раздел Schemas, где "
     "перечислены все компонентные схемы файла swagger.yaml. Раскрывая любую схему, "
     "можно увидеть тип каждого поля и формат (string, integer, number, date-time и т. д.), "
     "какие поля обязательны (required), а также примеры значений (example).")
body(doc, "Это особенно полезно при разработке клиентского приложения: разработчик видит "
     "точную структуру объектов без необходимости читать исходный код базы данных.")
body(doc, "Основные схемы API приведены в таблице ниже. Схема Client включает поля id, "
     "name, email, phone, city, registered_at. Схема ClientInput — name (обязательное), "
     "email (обязательное), phone, city. Схема Item — id, name, category, price, stock. "
     "Схема Order — id, client_id, item_id, status_id, quantity, total. Схемы "
     "LoginRequest и LoginResponse описывают запрос и ответ авторизации.")

embed_image(doc, 18, "Раздел Schemas: развёрнутые схемы Client и ClientInput с типами полей.")

# ─── РАЗДЕЛ 9 ─────────────────────────────────────────────────────────────────
heading(doc, "Вывод", 9)

body(doc, "В ходе выполнения дополнительного раздела работы была разработана спецификация "
     "REST API в формате OpenAPI 3.0, которая охватывает все четыре таблицы базы данных "
     "и полностью описывает операции авторизации, получения, добавления, изменения и "
     "удаления записей. Спецификация хранится в файле docs/swagger.yaml и содержит "
     "переиспользуемые схемы данных, параметры запросов и настройку JWT-авторизации.")
body(doc, "Визуализация выполнена через Swagger UI: файл docs/index.html без каких-либо "
     "фреймворков или сборщиков отображает интерактивную документацию. Все основные "
     "операции были проверены через кнопку Try it out непосредственно в браузере — "
     "POST, PATCH, DELETE с авторизацией по Bearer Token и следующим за ними "
     "проверочным GET, что полностью соответствует сценариям, отработанным в Postman.")
body(doc, "Использование Swagger/OpenAPI упрощает взаимодействие между командой разработки "
     "и тестировщиками: актуальная документация всегда доступна в браузере и не требует "
     "синхронизации вручную.")

# ─── СОХРАНЕНИЕ ───────────────────────────────────────────────────────────────
doc.save(OUT_PATH)
print(f"\nФайл сохранён: {OUT_PATH}")
